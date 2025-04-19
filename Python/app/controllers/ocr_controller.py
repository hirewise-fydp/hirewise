from flask import request, jsonify
from models.pdf_utils import is_pdf_scanned, process_scanned_pdf
from models.ocr_utils import perform_ocr_on_image
import os
import fitz
import requests
import tempfile
import time
from docx import Document

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'Uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)  # Ensure the uploads folder exists

def extract_text():
    try:
        # Expecting a JSON payload with an image URL
        data = request.get_json()
        if not data or 'image_url' not in data:
            return jsonify({'error': 'No image URL provided'}), 400

        image_url = data['image_url']
        print(f'Processing image URL: {image_url}')

        # Download the image with retries
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        max_retries = 3
        retry_delay = 1  # seconds
        for attempt in range(1, max_retries + 1):
            try:
                response = requests.get(image_url, headers=headers, stream=True, timeout=30)
                print(f'Attempt {attempt}: Response status: {response.status_code}, Content-Type: {response.headers.get("content-type", "")}')
                if response.status_code == 200:
                    break
                else:
                    print(f'Attempt {attempt}: Non-200 status code: {response.status_code}, Response: {response.text}')
                    if attempt == max_retries:
                        return jsonify({'error': f'Failed to download image from URL after {max_retries} attempts, status code: {response.status_code}'}), 400
            except requests.exceptions.RequestException as e:
                print(f'Attempt {attempt}: Failed to download image from {image_url}: {str(e)}')
                if attempt == max_retries:
                    return jsonify({'error': f'Failed to download image from URL after {max_retries} attempts: {str(e)}'}), 400
            time.sleep(retry_delay * attempt)

        # Determine file extension from URL, Content-Type, or Content-Disposition
        content_type = response.headers.get('content-type', '').lower()
        content_disposition = response.headers.get('content-disposition', '').lower()
        url_ext = os.path.splitext(image_url)[1].lower()

        # Extract filename from Content-Disposition if available
        filename = None
        if 'filename=' in content_disposition:
            filename = content_disposition.split('filename=')[-1].strip('"\'')
            filename_ext = os.path.splitext(filename)[1].lower()

        if 'pdf' in content_type or url_ext == '.pdf' or (filename and filename_ext == '.pdf'):
            file_ext = '.pdf'
        elif 'png' in content_type or url_ext == '.png' or (filename and filename_ext == '.png'):
            file_ext = '.png'
        elif 'docx' in content_type or url_ext == '.docx' or (filename and filename_ext == '.docx'):
            file_ext = '.docx'
        elif 'jpeg' in content_type or 'jpg' in content_type or url_ext in ('.jpg', '.jpeg') or (filename and filename_ext in ('.jpg', '.jpeg')):
            file_ext = '.jpg'
        else:
            print(f'Unsupported content-type: {content_type}, URL extension: {url_ext}, Content-Disposition filename: {filename}')
            return jsonify({'error': 'Unsupported file format'}), 400

        # Create a temporary file to store the downloaded content
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_ext, dir=UPLOAD_FOLDER)
        input_path = temp_file.name
        with open(input_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        print(f'Saved temporary file: {input_path}')

        extracted_text = ""

        if file_ext == '.pdf':
            if is_pdf_scanned(input_path):
                extracted_text = process_scanned_pdf(input_path)
            else:
                pdf_document = fitz.open(input_path)
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    extracted_text += page.get_text("text") + "\n"
                pdf_document.close()
        elif file_ext in ('.png', '.jpg', '.jpeg'):
            extracted_text = perform_ocr_on_image(input_path)
        elif file_ext == '.docx':
            doc = Document(input_path)
            extracted_text = ""

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    extracted_text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():  # Only include non-empty cells
                            extracted_text += cell.text + "\n"

            # Extract text from headers and footers
            for section in doc.sections:
                # Headers
                for header in section.header.paragraphs:
                    if header.text.strip():
                        extracted_text += header.text + "\n"
                # Footers
                for footer in section.footer.paragraphs:
                    if footer.text.strip():
                        extracted_text += footer.text + "\n"
        else:
            os.remove(input_path)
            return jsonify({'error': 'Unsupported file format'}), 400

        # Clean up the temporary file
        try:
            os.remove(input_path)
        except Exception as e:
            print(f'Error cleaning up temporary file {input_path}: {e}')

        return jsonify({'text': extracted_text.strip()})

    except Exception as e:
        print(f'Error in extract_text: {e}')
        # Ensure temporary file is cleaned up in case of error
        if 'input_path' in locals():
            try:
                os.remove(input_path)
            except Exception:
                pass
        return jsonify({'error': str(e)}), 500